from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file
from werkzeug.security import generate_password_hash, check_password_hash
from functools import wraps
from datetime import datetime
import sqlite3
import os
from reportlab.pdfgen import canvas
from docx import Document

app = Flask(__name__)
app.secret_key = "secret_key_here"

DATABASE = "notebook.db"


# ---------------- DATABASE HELPER ---------------- #

def get_db():
    conn = sqlite3.connect(DATABASE)
    conn.row_factory = sqlite3.Row
    return conn


# ---------------- LOGIN REQUIRED DECORATOR ---------------- #

def login_required(route_function):
    @wraps(route_function)
    def wrapper(*args, **kwargs):
        if "user_id" not in session:
            return redirect(url_for("login"))
        return route_function(*args, **kwargs)
    return wrapper


# ---------------- HOME ---------------- #

@app.route("/")
def home():
    if "user_id" in session:
        return redirect(url_for("show_notes"))
    return redirect(url_for("login"))


# ---------------- SIGNUP ---------------- #

@app.route("/signup", methods=["GET", "POST"])
def signup():

    if request.method == "POST":

        name = request.form.get("name")
        email = request.form.get("email")
        password = request.form.get("password")
        dob = request.form.get("dob")

        if not name or not email or not password:
            flash("All fields are required")
            return redirect(url_for("signup"))

        hashed_password = generate_password_hash(password)

        db = get_db()
        cursor = db.cursor()

        try:
            cursor.execute(
                "INSERT INTO user(name,email,password,dob) VALUES(?,?,?,?)",
                (name, email, hashed_password, dob)
            )
            db.commit()

            flash("Account created successfully")
            return redirect(url_for("login"))

        except sqlite3.IntegrityError:
            flash("Email already exists")

        finally:
            db.close()

    return render_template("signup.html")


# ---------------- LOGIN ---------------- #

@app.route("/login", methods=["GET", "POST"])
def login():

    if request.method == "POST":

        email = request.form.get("email")
        password = request.form.get("password")

        db = get_db()

        user = db.execute(
            "SELECT * FROM user WHERE email=?",
            (email,)
        ).fetchone()

        db.close()

        if user and check_password_hash(user["password"], password):

            if user["is_deleted"] == 1:
                flash("Account scheduled for deletion")
                return redirect(url_for("login"))

            session["user_id"] = user["id"]
            return redirect(url_for("profile"))

        flash("Invalid email or password")

    return render_template("login.html")


# ---------------- LOGOUT ---------------- #

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


# ---------------- PROFILE ---------------- #

@app.route("/profile")
@login_required
def profile():

    db = get_db()

    user = db.execute(
        "SELECT * FROM user WHERE id=?",
        (session["user_id"],)
    ).fetchone()

    db.close()

    return render_template("profile.html", user=user)


# ---------------- CREATE NOTE ---------------- #

@app.route("/create_note", methods=["GET", "POST"])
@login_required
def create_note():

    if request.method == "POST":

        title = request.form.get("title")
        content = request.form.get("content")

        db = get_db()

        db.execute(
            "INSERT INTO notebook(title,content,user_id,created_at) VALUES(?,?,?,?)",
            (title, content, session["user_id"], datetime.now())
        )

        db.commit()
        db.close()

        return redirect(url_for("show_notes"))

    return render_template("create_note.html")


# ---------------- SHOW NOTES ---------------- #

@app.route("/notes")
@login_required
def show_notes():

    db = get_db()

    notes = db.execute(
        "SELECT * FROM notebook WHERE user_id=? AND is_deleted=0",
        (session["user_id"],)
    ).fetchall()

    db.close()

    return render_template("notes.html", notes=notes)


# ---------------- UPDATE NOTE ---------------- #

@app.route("/update_note/<int:note_id>", methods=["GET", "POST"])
@login_required
def update_note(note_id):

    db = get_db()

    if request.method == "POST":

        title = request.form.get("title")
        content = request.form.get("content")

        db.execute(
            "UPDATE notebook SET title=?,content=? WHERE id=? AND user_id=?",
            (title, content, note_id, session["user_id"])
        )

        db.commit()
        db.close()

        return redirect(url_for("show_notes"))

    note = db.execute(
        "SELECT * FROM notebook WHERE id=? AND user_id=?",
        (note_id, session["user_id"])
    ).fetchone()

    db.close()

    return render_template("update_note.html", note=note)


# ---------------- MOVE TO TRASH ---------------- #

@app.route("/move_to_trash/<int:note_id>", methods=["POST"])
@login_required
def move_to_trash(note_id):

    db = get_db()

    db.execute(
        "UPDATE notebook SET is_deleted=1 WHERE id=? AND user_id=?",
        (note_id, session["user_id"])
    )

    db.commit()
    db.close()

    return redirect(url_for("show_notes"))


# ---------------- TRASH PAGE ---------------- #

@app.route("/trash")
@login_required
def trash():

    db = get_db()

    notes = db.execute(
        "SELECT * FROM notebook WHERE user_id=? AND is_deleted=1",
        (session["user_id"],)
    ).fetchall()

    db.close()

    return render_template("trash.html", trash=notes)


# ---------------- RESTORE NOTE ---------------- #

@app.route("/restore_note/<int:note_id>", methods=["POST"])
@login_required
def restore_note(note_id):

    db = get_db()

    db.execute(
        "UPDATE notebook SET is_deleted=0 WHERE id=? AND user_id=?",
        (note_id, session["user_id"])
    )

    db.commit()
    db.close()

    return redirect(url_for("trash"))


# ---------------- SEARCH ---------------- #

@app.route("/search")
@login_required
def search():

    query = request.args.get("q", "")

    db = get_db()

    notes = db.execute(
        """
        SELECT * FROM notebook
        WHERE user_id=? AND is_deleted=0
        AND (title LIKE ? OR content LIKE ?)
        """,
        (session["user_id"], f"%{query}%", f"%{query}%")
    ).fetchall()

    db.close()

    return render_template("notes.html", notes=notes, query=query)


# ---------------- FILTER NOTES ---------------- #

@app.route("/filter")
@login_required
def filter_notes():

    method = request.args.get("method", "latest")

    order_map = {
        "latest": "created_at DESC",
        "oldest": "created_at ASC",
        "title": "title ASC"
    }

    order_by = order_map.get(method, "created_at DESC")

    db = get_db()

    notes = db.execute(
        f"SELECT * FROM notebook WHERE user_id=? AND is_deleted=0 ORDER BY {order_by}",
        (session["user_id"],)
    ).fetchall()

    db.close()

    return render_template("notes.html", notes=notes)


# ---------------- DELETE ACCOUNT ---------------- #

@app.route("/delete_account", methods=["POST"])
@login_required
def delete_account():

    db = get_db()

    db.execute(
        "UPDATE user SET is_deleted=1, delete_request_at=? WHERE id=?",
        (datetime.now(), session["user_id"])
    )

    db.commit()
    db.close()

    session.clear()

    return redirect(url_for("login"))


# ---------------- RECOVER ACCOUNT ---------------- #

@app.route("/recover_account", methods=["GET", "POST"])
def recover_account():

    if request.method == "POST":

        email = request.form.get("email")
        password = request.form.get("password")

        db = get_db()

        user = db.execute(
            "SELECT * FROM user WHERE email=?",
            (email,)
        ).fetchone()

        if not user:
            flash("User not found")

        elif not check_password_hash(user["password"], password):
            flash("Wrong password")

        else:
            db.execute(
                "UPDATE user SET is_deleted=0, delete_request_at=NULL WHERE email=?",
                (email,)
            )

            db.commit()

            flash("Account recovered successfully")
            return redirect(url_for("login"))

        db.close()

    return render_template("recover.html")


# ---------------- EXPORT NOTE ---------------- #

@app.route("/export_note/<int:note_id>")
@login_required
def export_note(note_id):

    file_type = request.args.get("type")

    db = get_db()

    note = db.execute(
        "SELECT * FROM notebook WHERE id=? AND user_id=?",
        (note_id, session["user_id"])
    ).fetchone()

    db.close()

    if not note:
        flash("Note not found")
        return redirect(url_for("show_notes"))

    title = note["title"]
    content = note["content"]

    file_path = f"{title}.{file_type}"

    if file_type == "pdf":

        c = canvas.Canvas(file_path)
        c.drawString(100, 800, title)
        c.drawString(100, 760, content)
        c.save()

    elif file_type == "docx":

        doc = Document()
        doc.add_heading(title)
        doc.add_paragraph(content)
        doc.save(file_path)

    return send_file(file_path, as_attachment=True)


# ---------------- RUN APP ---------------- #

if __name__ == "__main__":
    app.run(debug=True)