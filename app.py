from datetime import datetime, timedelta
from flask import Flask, jsonify, request, session,send_file
from reportlab.pdfgen import canvas
from docx import Document
import os.path
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from werkzeug.security import generate_password_hash, check_password_hash
from apscheduler.schedulers.background import BackgroundScheduler
import sqlite3
from werkzeug.utils import secure_filename
import pickle
from flask_cors import CORS
from flask_jwt_extended import (
JWTManager,
create_access_token,
create_refresh_token,
jwt_required,
get_jwt_identity
)
UPLOAD_FOLDER = 'uploads'
app = Flask(__name__)
CORS(app)
app.config["JWT_ACCESS_TOKEN_EXPIRES"] = timedelta(minutes=15)
app.config["JWT_REFRESH_TOKEN_EXPIRES"] = timedelta(days=7)
app.config["JWT_TOKEN_LOCATION"] = ["headers"]   # 🔥 IMPORTANT
app.config["JWT_HEADER_TYPE"] = "Bearer"
app.config["JWT_SECRET_KEY"] = "ashokkumaryadav"
jwt=JWTManager(app)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

def init_db():
    conn = sqlite3.connect('notebook.db')
    cur = conn.cursor()

    cur.execute('''
    CREATE TABLE IF NOT EXISTS user(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT,
        email TEXT UNIQUE,
        password TEXT,
        dob TEXT,
        created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
        delete_request_at DATETIME,
        is_deleted INTEGER DEFAULT 0
    )
    ''')

    cur.execute('''
    CREATE TABLE IF NOT EXISTS notebook(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT DEFAULT 'new_notebook',
        content TEXT,
        created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
        user_id INTEGER,
        is_deleted INTEGER DEFAULT 0,
        FOREIGN KEY(user_id) REFERENCES user(id) ON DELETE CASCADE
    )
    ''')
    cur.execute('''CREATE TABLE IF NOT EXISTS user_profile(id INTEGER  PRIMARY KEY AUTOINCREMENT, user_id INTEGER UNIQUE, dob TEXT, mobile TEXT, photo_path TEXT, secret_key TEXT, bio TEXT, FOREIGN KEY(user_id) REFERENCES user(id) ON DELETE CASCADE )''')
    #cur.execute('ALTER TABLE notebook ADD COLUMN category TEXT DEFAULT NULL')
    #cur.execute('ALTER TABLE notebook ADD COLUMN category TEXT DEFAULT NULL')

    conn.commit()
    conn.close()


init_db()


# ---------------- HOME ---------------- #


@app.route('/profile', methods=['GET'])
@jwt_required()
def profile():
    current_user=get_jwt_identity()
    # 2️⃣ Connect to DB
    conn = sqlite3.connect('notebook.db')
    conn.row_factory = sqlite3.Row  # So we can convert row to dict
    cur = conn.cursor()

    # 3️⃣ Fetch user data
    user = cur.execute(
        "SELECT id, name, email, dob FROM user WHERE id=?",
        (current_user,)
    ).fetchone()

    conn.close()

    if not user:
        return jsonify({"message": "User not found"}), 404

    return jsonify(dict(user))


# ---------------- SIGNUP ---------------- #

@app.route('/signup', methods=['POST'])
def signup():
    data = request.get_json()

    name = data.get('name')
    email = data.get('email')
    password = data.get('password')
    dob = data.get('dob')

    if not name or not email or not password:
        return jsonify({"message": "All fields required"}), 400

    hashed = generate_password_hash(password)

    conn = sqlite3.connect('notebook.db')
    cur = conn.cursor()

    try:
        cur.execute(
            "INSERT INTO user(name,email,password,dob) VALUES(?,?,?,?)",
            (name, email, hashed, dob)
        )
        conn.commit()
    except:
        return jsonify({"message": "Email already exists"}), 400

    conn.close()

    return jsonify({"message": "User created"})


# ---------------- LOGIN ---------------- #

@app.route('/login', methods=['POST'])
def login():

    data = request.get_json()

    email = data.get('email')
    password = data.get('password')

    conn = sqlite3.connect('notebook.db')
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()

    user = cur.execute(
        "SELECT * FROM user WHERE email=?",
        (email,)
    ).fetchone()

    if not user:
        return jsonify({"message": "User not found"}), 404

    if not check_password_hash(user['password'], password):
        return jsonify({"message": "Wrong password"}), 401

    if user['is_deleted'] == 1:
        return jsonify({"message": "Account scheduled for deletion. Recover within 30 days."})

    access_token = create_access_token(identity=str(user['id']))
    refresh_token = create_refresh_token(identity=str(user['id']))


    conn.close()

    return jsonify({"message": "Login successful", "access_token": access_token, "refresh_token": refresh_token})


@app.route('/refresh', methods=['POST'])
@jwt_required(refresh=True)
def refresh():
    current_user = get_jwt_identity()

    new_access_token = create_access_token(identity=current_user)

    return jsonify({
        "access_token": new_access_token
    })


# ---------------- CREATE NOTE ---------------- #

@app.route('/create_note', methods=['POST'])
@jwt_required()
def create_note():

    data = request.get_json()

    title = data.get('title')
    content = data.get('content')

    conn = sqlite3.connect('notebook.db')
    cur = conn.cursor()

    cur.execute(
        "INSERT INTO notebook(title,content,user_id) VALUES(?,?,?)",
        (title, content, get_jwt_identity())
    )

    conn.commit()
    conn.close()

    return jsonify({"message": "Note created"})


# ---------------- SHOW NOTES ---------------- #

@app.route('/notes', methods=['GET'])
@jwt_required()
def show_notes():

    conn = sqlite3.connect('notebook.db')
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()

    notes = cur.execute(
        "SELECT * FROM notebook WHERE user_id=? AND is_deleted=0",
        (get_jwt_identity(),)
    ).fetchall()

    conn.close()

    data = [dict(row) for row in notes]

    return jsonify({"notes": data})


# ---------------- UPDATE NOTE ---------------- #

@app.route('/update_notes/<int:id>', methods=['PUT'])
@jwt_required()
def update_note(id):

    data = request.get_json()

    title = data.get('title')
    content = data.get('content')

    conn = sqlite3.connect('notebook.db')
    cur = conn.cursor()

    cur.execute(
        "UPDATE notebook SET title=?,content=? WHERE id=? AND user_id=?",
        (title, content, id,get_jwt_identity())
    )

    conn.commit()
    conn.close()

    return jsonify({"message": "Note updated"})


# ---------------- MOVE TO TRASH ---------------- #

@app.route('/move_to_trash/<int:id>', methods=['PUT'])
@jwt_required()
def move_to_trash(id):

    conn = sqlite3.connect('notebook.db')
    cur = conn.cursor()

    cur.execute(
        "UPDATE notebook SET is_deleted=1 WHERE id=? AND user_id=?",
        (id, get_jwt_identity())
    )

    conn.commit()
    conn.close()

    return jsonify({"message": "Moved to trash"})


# ---------------- TRASH NOTES ---------------- #

@app.route('/trash', methods=['GET'])
@jwt_required()
def trash():

    conn = sqlite3.connect('notebook.db')
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()

    notes = cur.execute(
        "SELECT * FROM notebook WHERE user_id=? AND is_deleted=1",
        (get_jwt_identity(),)
    ).fetchall()

    conn.close()

    data = [dict(row) for row in notes]

    return jsonify({"trash": data})


# ---------------- RESTORE NOTE ---------------- #

@app.route('/restore_note/<int:id>', methods=['PUT'])
@jwt_required()
def restore_note(id):

    conn = sqlite3.connect('notebook.db')
    cur = conn.cursor()

    cur.execute(
        "UPDATE notebook SET is_deleted=0 WHERE id=? AND user_id=?",
        (id, get_jwt_identity())
    )

    conn.commit()
    conn.close()

    return jsonify({"message": "Note restored"})


# ---------------- SEARCH ---------------- #

@app.route('/search', methods=['POST'])
@jwt_required()
def search():
    query = request.json.get('query')

    conn = sqlite3.connect('notebook.db')
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()

    notes = cur.execute("""
    SELECT * FROM notebook
    WHERE user_id=? AND is_deleted=0
    AND (title LIKE ? OR content LIKE ?)
    """, (
        get_jwt_identity(),
        '%' + query + '%',
        '%' + query + '%'
    )).fetchall()

    conn.close()

    result = [dict(row) for row in notes]

    return jsonify({"results": result})


# ---------------- FILTER ---------------- #

@app.route('/filter', methods=['POST'])
@jwt_required()
def filter_notes():

    method = request.json.get('method')

    conn = sqlite3.connect('notebook.db')
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()

    if method == "latest":
        query = "ORDER BY created_at DESC"

    elif method == "oldest":
        query = "ORDER BY created_at ASC"

    elif method == "title":
        query = "ORDER BY title ASC"

    else:
        return jsonify({"message": "Invalid method"})

    notes = cur.execute(
        f"SELECT * FROM notebook WHERE user_id=? AND is_deleted=0 {query}",
        (get_jwt_identity(),)
    ).fetchall()

    conn.close()

    result = [dict(row) for row in notes]

    return jsonify({"notes": result})


# ---------------- DELETE ACCOUNT REQUEST ---------------- #

@app.route('/delete_account', methods=['POST'])
@jwt_required()
def delete_account():
    conn = sqlite3.connect('notebook.db')
    cur = conn.cursor()
    cur.execute(
        "UPDATE user SET is_deleted=1,delete_request_at=? WHERE id=?",
        (datetime.now(), get_jwt_identity())
    )

    conn.commit()
    conn.close()

    session.clear()

    return jsonify({"message": "Account will be deleted after 30 days"})


# ---------------- RECOVER ACCOUNT ---------------- #

@app.route('/recover_account', methods=['POST'])
def recover_account():

    data = request.json
    email = data.get('email')
    password = data.get('password')

    conn = sqlite3.connect('notebook.db')
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()

    user = cur.execute(
        "SELECT * FROM user WHERE email=?",
        (email,)
    ).fetchone()

    if not user:
        return jsonify({"message": "User not found"})

    if not check_password_hash(user['password'], password):
        return jsonify({"message": "Wrong password"})

    cur.execute(
        "UPDATE user SET is_deleted=0,delete_request_at=NULL WHERE email=?",
        (email,)
    )

    conn.commit()
    conn.close()

    return jsonify({"message": "Account recovered"})


# ---------------- AUTO DELETE AFTER 30 DAYS ---------------- #

def delete_old_users():

    conn = sqlite3.connect('notebook.db')
    cur = conn.cursor()

    cur.execute("""
    DELETE FROM user
    WHERE is_deleted=1
    AND delete_request_at <= datetime('now','-30 days')
    """)

    conn.commit()
    conn.close()



@app.route('/export_note/<int:id>', methods=['GET'])
@jwt_required()
def export_note(id):
    user_id = get_jwt_identity()

    filetype = request.args.get("type")   # pdf / docx

    conn = sqlite3.connect('notebook.db')
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()

    note = cur.execute(
        "SELECT * FROM notebook WHERE id=? AND user_id=?",
        (id,user_id)
    ).fetchone()

    conn.close()

    if not note:
        return jsonify({"message":"note not found"})

    title = note['title']
    content = note['content']

    os.makedirs("exports", exist_ok=True)

    # -------- PDF EXPORT -------- #

    if filetype == "pdf":

        filepath = f"exports/{title}.pdf"

        c = canvas.Canvas(filepath)
        c.setFont("Helvetica",12)

        y = 800
        for line in content.split("\n"):
            c.drawString(50,y,line)
            y -= 20

        c.save()

        return send_file(filepath, as_attachment=True)


    # -------- DOCX EXPORT -------- #

    elif filetype == "docx":

        filepath = f"exports/{title}.docx"

        doc = Document()
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)

        doc.save(filepath)

        return send_file(filepath, as_attachment=True)

    else:
        return jsonify({"message":"invalid file type"})


@app.route('/share_note/<int:id>', methods=['POST'])
@jwt_required()
def share_note(id):

    user_id = get_jwt_identity()

    data = request.get_json()
    receiver_email = data.get('email')
    conn = sqlite3.connect('notebook.db')
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()

    note = cur.execute(
        "SELECT * FROM notebook WHERE id=? AND user_id=?",
        (id,user_id)
    ).fetchone()

    conn.close()

    if not note:
        return jsonify({"message":"note not found"})

    title = note['title']
    content = note['content']

    # -------- EMAIL CONFIG -------- #

    sender_email = "ashokjuriya3521@gmail.com"
    sender_password = "absbdfjadsfkjaddhfiajndnkinjfnfidfiuniuanfuiadfiaafi"

    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = receiver_email
    msg['Subject'] = f"Shared Note: {title}"

    body = f"""
Title: {title}

Content:
{content}
"""

    msg.attach(MIMEText(body,'plain'))

    try:

        server = smtplib.SMTP('smtp.gmail.com',587)
        server.starttls()
        server.login(sender_email,sender_password)

        server.send_message(msg)
        server.quit()

        return jsonify({"message":"note shared successfully via email"})

    except Exception as e:
        return jsonify({"error":str(e)})

@app.route('/update_user/<int:id>', methods=['PUT'])
@jwt_required()
def update_account(id):
    user_id=get_jwt_identity()
    data=request.get_json()
    name = data.get('name')
    email = data.get('email')
    dob = data.get('dob')
    mobile = data.get('mobile')
    username = data.get('username')
    secret_key = data.get('secret_key')

    photo = request.files.get('photo')

    if not os.path.exists(UPLOAD_FOLDER):
        os.mkdir(UPLOAD_FOLDER)

    photo_path = None

    if photo:
        filename = secure_filename(photo.filename)
        photo_path = os.path.join(UPLOAD_FOLDER, filename)
        photo.save(photo_path)

    conn = sqlite3.connect('notebook.db')
    cur = conn.cursor()

    # update user table
    cur.execute("""
    UPDATE user 
    SET name=?, email=?
    WHERE id=?
    """,(name,email,user_id))

    # check profile exist
    cur.execute("SELECT id FROM user_profile WHERE user_id=?", (user_id,))
    profile = cur.fetchone()

    if profile:
        # update profile
        cur.execute("""
        UPDATE user_profile
        SET dob=?, photo_path=?,mobile=?,username=?, secret_key=?
        WHERE user_id=?
        """,(dob,photo_path,mobile,username,secret_key,user_id))
    else:
        # insert profile
        cur.execute("""
        INSERT INTO user_profile (user_id,dob,photo_path,secret_key)
        VALUES (?,?,?,?)
        """,(user_id,dob,photo_path,secret_key))

    conn.commit()
    conn.close()

    return jsonify({"message":"updated successfully"})

@app.route('/update_password', methods=['PUT'])
@jwt_required()
def update_password():
    user_id = get_jwt_identity()

    data = request.get_json()
    email = data.get('email')
    old_password = data.get('old_password')
    new_password = data.get('new_password')

    conn = sqlite3.connect('notebook.db')
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()

    result = cur.execute(
        'SELECT * FROM user WHERE id=?',
        (user_id,)
    ).fetchone()

    if not result:
        return jsonify({"message": "user not found"})

    # email check
    if email != result['email']:
        return jsonify({"message": "wrong email"})

    # old password check
    if not check_password_hash(result['password'], old_password):
        return jsonify({"message": "wrong old password"})

    # new password hash
    hashed_password = generate_password_hash(new_password)

    cur.execute(
        'UPDATE user SET password=? WHERE id=?',
        (hashed_password, user_id)
    )

    conn.commit()
    conn.close()

    return jsonify({"message": "password updated successfully"})

@app.route('/profile_dashboard', methods=['GET'])
@jwt_required()
def profile_dashboard():
    user_id = get_jwt_identity()
    conn = sqlite3.connect('notebook.db')
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()

    result = cur.execute("""
        SELECT 
            user.id,
            user.email,
            user_profile.dob,
            user_profile.photo_path,
            user_profile.secret_key,
            user_profile.mobile,
            user.name
        FROM user
        LEFT JOIN user_profile
        ON user.id = user_profile.user_id
        WHERE user.id = ?
    """, (user_id,)).fetchone()

    conn.close()

    if not result:
        return jsonify({"message": "user not found"})

    return jsonify(dict(result))




# ---------------- Load models once ----------------
with open("Models/tfidf_vectorizer.pkl", "rb") as f:
    vectorizer = pickle.load(f)

with open("Models/category_model.pkl", "rb") as f:
    model = pickle.load(f)

with open("Models/label_encoder.pkl", "rb") as f:
    le = pickle.load(f)

# ---------------- Predict Function ---------------- #
def predict_category(text):
    vec = vectorizer.transform([text])
    pred = model.predict(vec)
    category = le.inverse_transform(pred)
    return category[0]

# ---------------- API Route ---------------- #
@app.route('/assign_category/<int:note_id>', methods=['PUT'])
@jwt_required()
# this route has one issue i have to add category checking if category assign previously so need to alert user during api call
def assign_category(note_id):

    user_id = get_jwt_identity()

    conn = sqlite3.connect('notebook.db')
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()

    # Fetch the note
    note = cur.execute(
        "SELECT title, content FROM notebook WHERE id=? AND user_id=?",
        (note_id, user_id)
    ).fetchone()

    if not note:
        conn.close()
        return jsonify({"message": "Note not found"}), 404

    # Combine title + content for prediction
    text = f"{note['title']} {note['content']}"
    category = predict_category(text)

    # Update the category in DB
    cur.execute(
        "UPDATE notebook SET category=? WHERE id=? AND user_id=?",
        (category, note_id, user_id)
    )
    conn.commit()
    conn.close()

    return jsonify({
        "message": "Category assigned",
        "category": category
    })

scheduler = BackgroundScheduler()
scheduler.add_job(delete_old_users, 'interval', hours=24)
scheduler.start()


# ---------------- RUN SERVER ---------------- #'''

if __name__ == "__main__":
    app.run(debug=True)

